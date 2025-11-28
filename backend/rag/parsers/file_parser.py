import io
import json
from typing import Optional, List, Dict, Any
import pandas as pd

class FileParser:
    """다양한 파일 형식을 텍스트로 변환하는 파서"""
    
    @staticmethod
    def parse_file(filename: str, file_content: bytes) -> str:
        """파일을 텍스트로 변환"""
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if file_ext == 'json':
            # JSON 파일은 구조화된 파서로 처리 (parse_json_file 사용)
            # 여기서는 기본 텍스트 변환만 수행
            try:
                decoded = file_content.decode('utf-8')
                return decoded
            except UnicodeDecodeError:
                try:
                    return file_content.decode('utf-8-sig')
                except UnicodeDecodeError:
                    return file_content.decode('latin-1')
        elif file_ext in ['xlsx', 'xls']:
            return FileParser._parse_excel(file_content)
        elif file_ext == 'csv':
            return FileParser._parse_csv(file_content, filename)
        elif file_ext in ['docx']:
            return FileParser._parse_docx(file_content)
        elif file_ext in ['doc']:
            return FileParser._parse_doc(file_content, filename)
        elif file_ext in ['txt', 'md', 'text']:
            return file_content.decode('utf-8')
        else:
            # 기본적으로 텍스트로 처리
            try:
                decoded = file_content.decode('utf-8')
                # HTML/JavaScript 코드인지 확인
                if '<html' in decoded.lower() or '<script' in decoded.lower():
                    # HTML에서 텍스트 추출 시도
                    return FileParser._extract_text_from_html(decoded)
                return decoded
            except UnicodeDecodeError:
                try:
                    return file_content.decode('utf-8-sig')  # BOM 제거
                except UnicodeDecodeError:
                    return file_content.decode('latin-1')  # 폴백
    
    @staticmethod
    def _parse_excel(file_content: bytes) -> str:
        """Excel 파일을 텍스트로 변환"""
        try:
            # 모든 시트 읽기
            excel_file = io.BytesIO(file_content)
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            
            text_parts = []
            for sheet_name, df in excel_data.items():
                text_parts.append(f"=== 시트: {sheet_name} ===")
                text_parts.append(FileParser._dataframe_to_text(df))
                text_parts.append("")  # 빈 줄 추가
            
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Excel 파일 파싱 오류: {str(e)}")
    
    @staticmethod
    def _parse_csv(file_content: bytes, filename: str) -> str:
        """CSV 파일을 텍스트로 변환"""
        try:
            # 인코딩 자동 감지 시도
            encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1']
            separators = [',', ';', '\t', '|']
            df = None
            
            for encoding in encodings:
                for sep in separators:
                    try:
                        csv_file = io.BytesIO(file_content)
                        df = pd.read_csv(
                            csv_file, 
                            encoding=encoding,
                            sep=sep,
                            on_bad_lines='skip'  # 잘못된 줄 건너뛰기
                        )
                        # 최소한의 데이터가 있는지 확인
                        if len(df) > 0 or len(df.columns) > 0:
                            break
                    except (UnicodeDecodeError, pd.errors.ParserError, ValueError):
                        continue
                if df is not None and (len(df) > 0 or len(df.columns) > 0):
                    break
            
            if df is None or (len(df) == 0 and len(df.columns) == 0):
                raise ValueError("CSV 파일을 파싱할 수 없습니다. 인코딩이나 구분자를 확인해주세요.")
            
            return FileParser._dataframe_to_text(df)
        except Exception as e:
            raise ValueError(f"CSV 파일 파싱 오류: {str(e)}")
    
    @staticmethod
    def _dataframe_to_text(df: pd.DataFrame) -> str:
        """DataFrame을 RAG에 적합한 텍스트로 변환"""
        text_parts = []
        
        # 컬럼명 추가
        columns = " | ".join([str(col) for col in df.columns])
        text_parts.append(f"컬럼: {columns}")
        text_parts.append("")
        
        # 각 행을 텍스트로 변환
        for idx, row in df.iterrows():
            row_text_parts = []
            for col in df.columns:
                value = row[col]
                if pd.notna(value):
                    row_text_parts.append(f"{col}: {value}")
                else:
                    row_text_parts.append(f"{col}: (비어있음)")
            
            row_text = " | ".join(row_text_parts)
            text_parts.append(f"행 {idx + 1}: {row_text}")
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _parse_docx(file_content: bytes) -> str:
        """DOCX 파일을 텍스트로 변환"""
        try:
            from docx import Document
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 테이블 처리
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError("python-docx 라이브러리가 필요합니다. pip install python-docx")
        except Exception as e:
            raise ValueError(f"DOCX 파일 파싱 오류: {str(e)}")
    
    @staticmethod
    def _parse_doc(file_content: bytes, filename: str) -> str:
        """DOC 파일을 텍스트로 변환"""
        try:
            # 먼저 docx2python 시도 (.doc 파일이 실제로는 .docx일 수 있음)
            try:
                from docx2python import docx2python
                doc_file = io.BytesIO(file_content)
                doc = docx2python(doc_file)
                text = doc.text
                if text and len(text.strip()) > 100:  # 충분한 텍스트가 추출되었는지 확인
                    return text
            except:
                pass
            
            # 텍스트로 디코딩 시도
            try:
                decoded = file_content.decode('utf-8')
                # HTML인 경우 텍스트 추출
                if '<html' in decoded.lower() or '<script' in decoded.lower():
                    return FileParser._extract_text_from_html(decoded)
                return decoded
            except:
                pass
            
            # 다른 인코딩 시도
            for encoding in ['cp949', 'euc-kr', 'latin-1']:
                try:
                    decoded = file_content.decode(encoding)
                    if '<html' in decoded.lower() or '<script' in decoded.lower():
                        return FileParser._extract_text_from_html(decoded)
                    return decoded
                except:
                    continue
            
            raise ValueError("DOC 파일을 파싱할 수 없습니다.")
        except Exception as e:
            raise ValueError(f"DOC 파일 파싱 오류: {str(e)}")
    
    @staticmethod
    def _extract_text_from_html(html_content: str) -> str:
        """HTML에서 텍스트 추출"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 스크립트와 스타일 제거
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 텍스트 추출
            text = soup.get_text()
            
            # 줄바꿈 정리
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except ImportError:
            # BeautifulSoup이 없으면 기본 정리만 수행
            import re
            # HTML 태그 제거
            text = re.sub(r'<[^>]+>', '', html_content)
            # 여러 공백을 하나로
            text = re.sub(r'\s+', ' ', text)
            return text.strip()
        except Exception as e:
            # 오류 발생 시 원본 반환
            return html_content
    
    @staticmethod
    def parse_law_table_file(file_content: bytes, filename: str, scenario: str) -> List[Dict[str, Any]]:
        """5-2.법령별 문서를 구조화된 청크로 파싱"""
        try:
            from rag.parsers.law_table_parser import LawTableParser
            
            # HTML로 디코딩
            html_content = file_content.decode('utf-8')
            
            # HTML이 아니면 시도
            if '<html' not in html_content.lower():
                # 다른 인코딩 시도
                for encoding in ['cp949', 'euc-kr']:
                    try:
                        html_content = file_content.decode(encoding)
                        if '<html' in html_content.lower() or '<table' in html_content.lower():
                            break
                    except:
                        continue
            
            # 법령 테이블 파싱
            chunks = LawTableParser.parse_law_table(
                html_content=html_content,
                scenario=scenario,
                filename=filename
            )
            
            return chunks
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"법령 테이블 파싱 오류: {str(e)}")
            # 실패 시 일반 텍스트 추출
            return []
    
    @staticmethod
    def parse_json_file(file_content: bytes, filename: str, folder: str = None) -> List[Dict[str, Any]]:
        """JSON 파일을 구조화된 청크로 파싱 (토큰 단위 처리)"""
        try:
            # JSON 디코딩
            json_content = file_content.decode('utf-8')
            data = json.loads(json_content)
            
            if not isinstance(data, list):
                # 배열이 아니면 단일 객체를 배열로 변환
                data = [data]
            
            chunks = []
            for item in data:
                if not isinstance(item, dict):
                    continue
                
                # JSON 항목을 구조화된 텍스트로 변환
                chunk_text_parts = []
                
                # Q&A 형식 (Construction_law_qa.json)
                if "question" in item and "answer" in item:
                    chunk_text_parts.append(f"질문: {item.get('question', '')}")
                    chunk_text_parts.append(f"답변: {item.get('answer', '')}")
                    
                    if item.get('category'):
                        chunk_text_parts.append(f"카테고리: {item.get('category')}")
                    if item.get('keywords'):
                        keywords = item.get('keywords', [])
                        if isinstance(keywords, list):
                            chunk_text_parts.append(f"키워드: {', '.join(keywords)}")
                    if item.get('legal_basis'):
                        legal_basis = item.get('legal_basis', [])
                        if isinstance(legal_basis, list) and legal_basis:
                            chunk_text_parts.append(f"법적 근거: {', '.join(str(b) for b in legal_basis)}")
                    if item.get('reference_document'):
                        chunk_text_parts.append(f"참고 문서: {item.get('reference_document')}")
                    if item.get('application_scope'):
                        chunk_text_parts.append(f"적용 범위: {item.get('application_scope')}")
                
                # 조례 형식 (Jeonju_Construction_Ordinance.json)
                elif "title" in item and "answer" in item:
                    chunk_text_parts.append(f"제목: {item.get('title', '')}")
                    if item.get('question'):
                        chunk_text_parts.append(f"질문: {item.get('question')}")
                    chunk_text_parts.append(f"답변: {item.get('answer', '')}")
                    
                    if item.get('category'):
                        chunk_text_parts.append(f"카테고리: {item.get('category')}")
                    if item.get('source_file'):
                        chunk_text_parts.append(f"출처 파일: {item.get('source_file')}")
                    if item.get('keywords'):
                        keywords = item.get('keywords', [])
                        if isinstance(keywords, list):
                            chunk_text_parts.append(f"키워드: {', '.join(keywords)}")
                    if item.get('regulation_type'):
                        chunk_text_parts.append(f"규정 유형: {item.get('regulation_type')}")
                    if item.get('jurisdiction'):
                        chunk_text_parts.append(f"관할: {item.get('jurisdiction')}")
                    if item.get('reference_law'):
                        chunk_text_parts.append(f"참고 법령: {item.get('reference_law')}")
                
                # 기본 형식 (다른 필드들)
                else:
                    # 모든 필드를 텍스트로 변환
                    for key, value in item.items():
                        if value is not None:
                            if isinstance(value, list):
                                chunk_text_parts.append(f"{key}: {', '.join(str(v) for v in value)}")
                            elif isinstance(value, dict):
                                chunk_text_parts.append(f"{key}: {json.dumps(value, ensure_ascii=False)}")
                            else:
                                chunk_text_parts.append(f"{key}: {value}")
                
                # 청크 텍스트 생성
                chunk_text = "\n".join(chunk_text_parts)
                
                # 메타데이터 구성
                chunk_metadata = {
                    "id": item.get("id", ""),
                    "filename": filename,
                    "folder": folder,
                    "source_file": item.get("source_file", ""),
                    "category": item.get("category", ""),
                    "json_type": "qa" if "question" in item else "ordinance" if "title" in item else "general"
                }
                
                # 원본 JSON 데이터도 메타데이터에 포함 (필요시 사용)
                chunk_metadata.update({k: v for k, v in item.items() if k not in chunk_metadata})
                
                chunks.append({
                    "content": chunk_text,
                    "metadata": chunk_metadata
                })
            
            return chunks
        except json.JSONDecodeError as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"JSON 파싱 오류: {filename}, {str(e)}")
            return []
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"JSON 파일 처리 오류: {filename}, {str(e)}")
            return []

